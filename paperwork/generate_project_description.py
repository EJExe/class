from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


OUTPUT_PATH = r"d:\DIPLOM\DIPLOM_LMS_vkr_report.docx"
ASSETS_ROOT = Path(r"d:\DIPLOM\doc_assets")
DIAGRAMS_DIR = ASSETS_ROOT / "diagrams"
SCREENSHOTS_DIR = ASSETS_ROOT / "screenshots"


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)


def set_paragraph_format(paragraph, first_line: float = 1.25, spacing: float = 1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(first_line)
    fmt.line_spacing = spacing
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(text)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.add_run(text)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_figure(document: Document, image_path: Path, caption: str, width_cm: float = 16.0) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(document, caption)


def add_code_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def read_code_excerpt(path: str, start_line: int, end_line: int) -> str:
    lines = Path(path).read_text(encoding="utf-8").splitlines()
    return "\n".join(lines[start_line - 1:end_line])


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    for line in code.splitlines():
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.add_break()


def add_cover_page(document: Document) -> None:
    lines = [
        "Пояснительная записка",
        "к описанию программного проекта",
        "",
        "Тема:",
        "«Разработка веб-приложения для управления учебными курсами,",
        "заданиями и проверкой работ»",
        "",
        "Проект: DIPLOM LMS",
        "",
        "Иваново, 2026",
    ]

    for index, line in enumerate(lines):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = paragraph.paragraph_format
        fmt.line_spacing = 1.5
        if index < 3:
            fmt.space_after = Pt(12)
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14 if index not in {0, 5} else 16)
        if index in {0, 5}:
            run.bold = True

    document.add_section(WD_SECTION.NEW_PAGE)


def add_keywords(document: Document) -> None:
    add_heading(document, "РЕФЕРАТ")
    add_body_paragraph(
        document,
        "Ключевые слова: LMS, учебный курс, учебные группы, assignment-канал, проверка работ, "
        "ведомость оценок, уведомления, аудит действий, WebSocket, WebRTC, локальное файловое хранилище.",
    )
    add_body_paragraph(
        document,
        "Объектом разработки является веб-приложение для сопровождения учебного процесса внутри курса: "
        "от регистрации участников и публикации материалов до сдачи, проверки и фиксации результатов.",
    )
    add_body_paragraph(
        document,
        "Цель проекта состоит в создании системы, которая объединяет в одном интерфейсе курсы, "
        "группы, каналы общения, задания, частную коммуникацию преподавателя со студентом, "
        "журнал оценивания и средства оперативного оповещения.",
    )
    add_body_paragraph(
        document,
        "Практический результат работы представляет собой локальную LMS-платформу на базе NestJS, Prisma, "
        "PostgreSQL и React. Реализация включает разграничение прав по ролям, чат с вложениями и реакциями, "
        "assignment-каналы, загрузку версий студенческих файлов, журнал проверки, экспорт ведомостей и видеокомнату.",
    )
    add_body_paragraph(
        document,
        "Область применения результатов охватывает учебные курсы кафедры, небольшие корпоративные программы "
        "обучения и внутренние образовательные проекты, где важны локальное хранение данных и прозрачная история работы.",
    )


def add_glossary(document: Document) -> None:
    add_heading(document, "ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ")
    glossary_items = [
        ("LMS", "Learning Management System, система управления обучением."),
        ("REST API", "набор HTTP-интерфейсов для обмена данными между клиентом и сервером."),
        ("WebSocket", "протокол двустороннего обмена данными в реальном времени."),
        ("WebRTC", "технология для организации аудио- и видеосвязи между клиентами."),
        ("CRUD", "базовые операции над данными: создание, чтение, изменение и удаление."),
        ("Assignment", "учебное задание, связанное с каналом курса и процессом сдачи работы."),
        ("Submission", "студенческая сдача по конкретному заданию, включая файлы, статус и комментарии."),
        ("Audit log", "журнал действий пользователей и системы."),
    ]

    for term, description in glossary_items:
        paragraph = document.add_paragraph()
        set_paragraph_format(paragraph, first_line=0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bold = paragraph.add_run(f"{term} ")
        bold.bold = True
        paragraph.add_run(description)


def add_roles_table(document: Document) -> None:
    add_heading(document, "1.4. Роли пользователей и заинтересованные лица", level=2)
    add_body_paragraph(
        document,
        "Система проектировалась не под абстрактного пользователя, а под несколько устойчивых ролей, "
        "которые по-разному работают с учебным контентом и по-разному отвечают за результат.",
    )
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Роль"
    header[1].text = "Основные действия"
    header[2].text = "Практический интерес"

    rows = [
        (
            "Администратор",
            "создание курсов, назначение ролей, экспорт данных, просмотр журнала аудита",
            "контроль структуры курса и прозрачность процессов",
        ),
        (
            "Преподаватель",
            "создание каналов и заданий, проверка работ, выставление оценок, общение со студентами",
            "оперативное сопровождение учебного процесса",
        ),
        (
            "Ассистент",
            "проверка работ, просмотр связанных материалов и личных чатов по заданиям",
            "поддержка преподавателя без доступа к лишним настройкам",
        ),
        (
            "Студент",
            "вступление в курс, работа в доступных каналах, загрузка черновиков и итоговых файлов",
            "понятный путь от получения задания до обратной связи",
        ),
    ]

    for role, actions, interest in rows:
        cells = table.add_row().cells
        cells[0].text = role
        cells[1].text = actions
        cells[2].text = interest


def add_entities_table(document: Document) -> None:
    add_heading(document, "2.3. Основные сущности данных", level=2)
    add_body_paragraph(
        document,
        "При проектировании схемы БД упор делался на то, чтобы данные отражали реальный ход работы, "
        "а не только конечный результат. Поэтому в модели есть не только курс и задание, но и состояния чтения, "
        "версии файлов, личные чаты, реакции и отдельные журналы активности.",
    )
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Сущность"
    header[1].text = "Назначение"
    header[2].text = "Ключевые поля"

    rows = [
        ("User", "хранение учетной записи и профиля", "login, email, nickname, avatarPath"),
        ("Course", "описание учебного курса", "title, description, inviteCode"),
        ("CourseMember", "роль пользователя внутри курса", "courseId, userId, role"),
        ("CourseGroup", "учебная группа внутри курса", "courseId, name"),
        ("Channel", "пространство общения или задание", "name, type, group access"),
        ("Assignment", "метаданные задания", "title, status, deadlineAt"),
        ("Submission", "студенческая сдача", "status, grade, teacherComment"),
        ("SubmissionFile", "отдельная версия файла", "originalName, uploadedAt"),
        ("Notification", "персональное уведомление", "type, title, isRead"),
        ("AuditLog", "журнал действий", "actionType, entityType, createdAt"),
    ]

    for entity, purpose, fields in rows:
        cells = table.add_row().cells
        cells[0].text = entity
        cells[1].text = purpose
        cells[2].text = fields


def add_testing_table(document: Document) -> None:
    add_heading(document, "4.1. Контрольные сценарии", level=2)
    add_body_paragraph(
        document,
        "Проверка системы строилась вокруг прикладных сценариев. Такой подход оказался полезнее, "
        "чем набор изолированных функций, потому что многие возможности завязаны друг на друга: "
        "например, сдача работы тянет за собой уведомления, обновление очереди проверки и изменения в истории задания.",
    )
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "№"
    header[1].text = "Сценарий"
    header[2].text = "Ожидаемый результат"
    header[3].text = "Статус"

    rows = [
        ("1", "Регистрация нового пользователя", "создается учетная запись и открывается рабочее пространство", "успешно"),
        ("2", "Создание курса преподавателем", "появляется курс с кодом приглашения и видеокомнатой", "успешно"),
        ("3", "Вступление студента по invite-коду", "пользователь добавляется в курс с ролью student", "успешно"),
        ("4", "Создание assignment-канала с дедлайном", "формируется канал и связанное задание", "успешно"),
        ("5", "Загрузка черновика студентом", "создается версия файла и запись в журнале активности", "успешно"),
        ("6", "Итоговая сдача работы", "статус меняется на submitted или submitted_late", "успешно"),
        ("7", "Проверка и выставление оценки", "студент получает комментарий и уведомление", "успешно"),
        ("8", "Отправка сообщения в курсовой чат", "сообщение появляется у участников без перезагрузки страницы", "успешно"),
        ("9", "Фильтрация файлов по курсу", "пользователь видит только доступные ему материалы", "успешно"),
        ("10", "Подключение к видеокомнате", "участник попадает в комнату при наличии свободного слота", "успешно"),
    ]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_sources_hint(document: Document) -> None:
    add_heading(document, "Заключение")
    add_body_paragraph(
        document,
        "В результате работы было подготовлено и реализовано веб-приложение, которое покрывает "
        "основные этапы учебного взаимодействия внутри курса. В проекте удалось совместить организацию "
        "учебных материалов, групповую коммуникацию, контроль сроков, проверку работ и фиксацию действий в одном контуре.",
    )
    add_body_paragraph(
        document,
        "Практическая ценность проекта определяется тем, что он не сводится к демонстрационному интерфейсу. "
        "Система опирается на реальную структуру данных, содержит разграничение прав, историю работы с файлами, "
        "экспортные сценарии и механизмы оперативной связи. Это делает ее пригодной как основу для дальнейшего "
        "развития под нужды конкретной кафедры или локальной образовательной платформы.",
    )
    add_body_paragraph(
        document,
        "В качестве следующего шага можно развивать календарное планирование, интеграцию с внешними "
        "источниками авторизации, автоматические напоминания по дедлайнам и более детальную аналитику по активности студентов. "
        "При этом уже в текущем виде проект демонстрирует целостный и технически обоснованный результат.",
    )


def add_appendix_two(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "ПРИЛОЖЕНИЕ 2")
    add_heading(document, "Фрагменты исходного кода программы", level=2)
    add_body_paragraph(
        document,
        "В приложение включены фрагменты исходного кода, которые характеризуют ключевые части проекта: "
        "описание структуры данных, бизнес-логику обработки сдач, realtime-взаимодействие и клиентские страницы.",
    )

    code_fragments = [
        (
            "Листинг 1 - Фрагмент схемы базы данных Prisma",
            read_code_excerpt(r"d:\DIPLOM\backend\prisma\schema.prisma", 93, 132),
        ),
        (
            "Листинг 2 - Обработка загрузки и отправки работы студентом",
            read_code_excerpt(r"d:\DIPLOM\backend\src\assignments\assignments.service.ts", 571, 742),
        ),
        (
            "Листинг 3 - События чата и видеокомнаты в realtime gateway",
            read_code_excerpt(r"d:\DIPLOM\backend\src\realtime\realtime.gateway.ts", 96, 239),
        ),
        (
            "Листинг 4 - Основная логика страницы курса",
            read_code_excerpt(r"d:\DIPLOM\frontend\src\pages\CoursePage.tsx", 30, 357),
        ),
        (
            "Листинг 5 - Основная логика страницы задания",
            read_code_excerpt(r"d:\DIPLOM\frontend\src\pages\AssignmentPage.tsx", 60, 352),
        ),
    ]

    for title, code in code_fragments:
        add_code_title(document, title)
        add_code_block(document, code)


def add_appendix_one(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "ПРИЛОЖЕНИЕ 1")
    add_heading(document, "Дополнительные диаграммы и экранные формы", level=2)
    add_body_paragraph(
        document,
        "В приложении приведены дополнительные иллюстрации, которые помогают быстро оценить композицию интерфейса "
        "и структуру программной системы.",
    )
    add_figure(document, DIAGRAMS_DIR / "architecture.png", "Рисунок П.1 - Архитектурная схема DIPLOM LMS", width_cm=16.0)
    add_figure(document, SCREENSHOTS_DIR / "course.png", "Рисунок П.2 - Основной экран курса", width_cm=16.0)
    add_figure(document, SCREENSHOTS_DIR / "assignment.png", "Рисунок П.3 - Экран работы с заданием", width_cm=14.2)


def fill_document(document: Document) -> None:
    add_cover_page(document)
    add_keywords(document)
    add_glossary(document)

    add_heading(document, "ВВЕДЕНИЕ")
    add_body_paragraph(
        document,
        "Учебный процесс чаще всего распадается на несколько независимых инструментов: отдельно ведется переписка, "
        "отдельно хранятся файлы, отдельно фиксируются оценки. На практике это создает лишние переходы между сервисами "
        "и заметно усложняет сопровождение курса. Преподавателю приходится вручную сверять сроки, версии файлов и "
        "комментарии, а студенту трудно быстро понять, где находится актуальное состояние его работы.",
    )
    add_body_paragraph(
        document,
        "Проект DIPLOM LMS вырос именно из этой проблемы. Его идея состоит не в том, чтобы повторить большую "
        "универсальную LMS, а в том, чтобы собрать компактную, понятную и локально развертываемую систему, "
        "в которой общение, задания и проверка связаны между собой напрямую.",
    )
    add_body_paragraph(
        document,
        "В качестве технологической основы выбрана клиент-серверная архитектура. Серверная часть реализована на NestJS "
        "с использованием Prisma и PostgreSQL, клиентская часть построена на React и Vite. Для сценариев, где задержка "
        "особенно заметна для пользователя, используются WebSocket и WebRTC.",
    )

    add_heading(document, "1. РАЗРАБОТКА И АНАЛИЗ ТРЕБОВАНИЙ")
    add_heading(document, "1.1. Актуальность темы", level=2)
    add_body_paragraph(
        document,
        "Для небольших учебных коллективов типична ситуация, когда преподаватель выдает задание в одном чате, "
        "материалы хранит в другой системе, а результаты проверки отправляет личными сообщениями. Пока группа небольшая, "
        "это кажется терпимым. Но как только в курсе появляются подгруппы, разные роли и несколько параллельных заданий, "
        "такой способ работы начинает давать сбои.",
    )
    add_body_paragraph(
        document,
        "Особенно остро проблема проявляется там, где важна история взаимодействия: кто и когда загрузил файл, "
        "какая версия считается актуальной, был ли студент уведомлен о замечаниях, что именно изменилось после проверки. "
        "В проекте сделан акцент именно на таких деталях, потому что они напрямую влияют на прозрачность учебного процесса.",
    )

    add_heading(document, "1.2. Анализ проблемы", level=2)
    add_body_paragraph(
        document,
        "Анализ существующей практики показал несколько устойчивых затруднений. Во-первых, каналы коммуникации и "
        "проверка работ обычно разорваны. Во-вторых, доступ к информации редко учитывает внутренние группы курса. "
        "В-третьих, в типичном решении преподаватель видит только конечную сдачу, а не путь, которым студент к ней пришел.",
    )
    add_body_paragraph(
        document,
        "Отдельную сложность вызывает работа с обратной связью. Если комментарии к заданию, личный диалог со студентом "
        "и сама оценка хранятся в разных местах, система быстро перестает быть прозрачной. Поэтому в рамках проекта "
        "было принято решение объединить эти процессы вокруг сущности задания.",
    )

    add_heading(document, "1.3. Цель и задачи разработки", level=2)
    add_body_paragraph(
        document,
        "Цель проекта заключается в разработке веб-приложения, которое поддерживает полный цикл работы с учебным курсом: "
        "от регистрации участников и выдачи заданий до проверки, оценки и хранения истории действий.",
    )
    add_body_paragraph(document, "Для достижения поставленной цели были определены следующие задачи:")
    for item in [
        "реализовать регистрацию, вход и редактирование профиля пользователя;",
        "ввести модель ролей с различающимися полномочиями внутри курса;",
        "создать механизм курсов, групп и каналов с учетом группового доступа;",
        "связать assignment-каналы с дедлайнами, файлами и состояниями выполнения;",
        "обеспечить поддержку черновиков, итоговой сдачи и проверки работ;",
        "реализовать персональные уведомления, очередь проверки и журнал аудита;",
        "сохранить внутри курса видеокомнату для синхронной работы участников.",
    ]:
        add_bullet(document, item)

    add_roles_table(document)

    add_heading(document, "1.5. Возможности системы", level=2)
    add_body_paragraph(
        document,
        "Функциональность проекта строится вокруг повседневных операций, которые действительно выполняются внутри курса. "
        "Пользователь после авторизации попадает в пространство курсов, где может создать новый курс или вступить в уже "
        "существующий по коду приглашения. Внутри курса доступны каналы общения, assignment-каналы, список участников, "
        "группы, ведомость и отдельные сервисные страницы.",
    )
    for item in [
        "создание курса с автоматическим формированием invite-кода и видеокомнаты;",
        "назначение ролей admin, teacher, assistant и student;",
        "формирование учебных групп и ограничение каналов по списку групп;",
        "чат курса с вложениями, редактированием, удалением и реакциями;",
        "упоминания пользователей в сообщениях;",
        "создание задания прямо на уровне канала с дедлайном и статусом;",
        "загрузка материалов задания преподавателем;",
        "загрузка студентом нескольких версий файла и финальная отправка работы;",
        "проверка работы, комментарии, выставление статуса и оценки;",
        "личный чат преподавателя и студента по конкретному заданию;",
        "очередь проверки, библиотека файлов, журнал аудита и экспортные сценарии.",
    ]:
        add_bullet(document, item)

    add_heading(document, "1.6. Функциональные требования", level=2)
    functional_requirements = [
        "Система должна обеспечивать регистрацию и аутентификацию пользователей.",
        "Система должна поддерживать создание курсов и вступление в них по invite-коду.",
        "Система должна разграничивать доступ к операциям по ролям пользователя.",
        "Система должна поддерживать групповой доступ к каналам курса.",
        "Система должна позволять создавать текстовые и assignment-каналы.",
        "Система должна хранить материалы заданий и версии студенческих файлов.",
        "Система должна фиксировать статус сдачи и проверки работы.",
        "Система должна отправлять уведомления о значимых событиях.",
        "Система должна предоставлять журнал аудита для административного контроля.",
        "Система должна поддерживать видеокомнату курса в реальном времени.",
    ]
    for item in functional_requirements:
        add_number(document, item)

    add_heading(document, "1.7. Нефункциональные требования", level=2)
    nonfunctional_requirements = [
        "Интерфейс должен оставаться понятным при повседневной работе без дополнительного обучения.",
        "Данные по курсу должны храниться локально, без обязательной привязки к внешним облачным сервисам.",
        "Критичные сценарии общения и уведомлений должны работать с минимальной задержкой.",
        "Структура проекта должна позволять добавлять новые модули без полной переработки архитектуры.",
        "История действий пользователя должна быть доступна для последующего анализа и разбора спорных ситуаций.",
    ]
    for item in nonfunctional_requirements:
        add_bullet(document, item)

    add_heading(document, "1.8. Диаграмма вариантов использования", level=2)
    add_body_paragraph(
        document,
        "На рисунке 1 показаны ключевые роли системы и основные действия, которые выполняются внутри платформы. "
        "Схема подчеркивает, что администратор, преподаватель и студент работают в одном приложении, но используют разные подмножества возможностей.",
    )
    add_figure(document, DIAGRAMS_DIR / "use_cases.png", "Рисунок 1 - Диаграмма вариантов использования системы", width_cm=16.2)

    add_heading(document, "2. ПРОЕКТИРОВАНИЕ")
    add_heading(document, "2.1. Архитектурное решение", level=2)
    add_body_paragraph(
        document,
        "Проект выполнен в виде SPA-приложения с выделенной серверной частью. Клиент отвечает за маршрутизацию, "
        "отображение данных и интерфейсные сценарии. Сервер предоставляет REST API, обрабатывает бизнес-логику, "
        "управляет хранилищем файлов, формирует уведомления и открывает WebSocket-шлюз для событий в реальном времени.",
    )
    add_body_paragraph(
        document,
        "Такое разделение позволило сохранить достаточно простую структуру и в то же время не смешивать в одном слое "
        "интерфейс и предметную логику. На клиенте основными страницами выступают курсы, курс, задание, ведомость, "
        "очередь проверки, библиотека файлов, уведомления, профиль и журнал аудита.",
    )
    add_body_paragraph(
        document,
        "На рисунке 2 приведена укрупненная архитектурная схема, показывающая взаимодействие клиентской части, серверных модулей, БД, файлового хранилища и realtime-компонентов.",
    )
    add_figure(document, DIAGRAMS_DIR / "architecture.png", "Рисунок 2 - Архитектура программной системы", width_cm=16.2)

    add_heading(document, "2.2. Логика доступа", level=2)
    add_body_paragraph(
        document,
        "Ограничение доступа реализовано не только по роли, но и по принадлежности к учебной группе. "
        "Преподаватель и администратор имеют управленческие права, ассистент включен в сценарии проверки, "
        "а студент видит только доступные ему каналы и не может открыть черновое задание до публикации.",
    )
    add_body_paragraph(
        document,
        "Отдельно выделены состояния чтения каналов и заданий. Благодаря этому пользователь видит непрочитанные элементы, "
        "а интерфейс может показывать локальные индикаторы без грубого сравнения по последней дате обновления курса целиком.",
    )

    add_entities_table(document)
    add_body_paragraph(
        document,
        "На рисунке 3 представлена концептуальная модель данных, которая отражает центральные сущности проекта и основные связи между ними.",
    )
    add_figure(document, DIAGRAMS_DIR / "data_model.png", "Рисунок 3 - Концептуальная модель данных проекта", width_cm=16.2)

    add_heading(document, "2.4. Проектирование интерфейсов", level=2)
    add_body_paragraph(
        document,
        "Интерфейс приложения сделан модульным. Главная страница курсов совмещает быстрый вход в рабочее пространство, "
        "создание нового курса и переход к уведомлениям. Внутри курса пользователь видит список каналов и рабочую область: "
        "либо чат, либо задание. Для преподавателя рядом с этим доступны инструменты управления каналами, группами и экспортом.",
    )
    add_body_paragraph(
        document,
        "Страница задания устроена как отдельный рабочий узел. Студент видит описание, файлы преподавателя, собственную сдачу, "
        "список версий и комментарии к каждой версии. Преподаватель в той же точке системы получает список работ студентов, "
        "форму проверки, личный чат, историю действий и аудит изменения задания.",
    )

    add_heading(document, "2.5. Технические средства и технологии", level=2)
    for item in [
        "frontend: React 18, TypeScript, Vite, React Router;",
        "backend: NestJS, TypeScript, Socket.IO;",
        "доступ к данным: Prisma ORM;",
        "СУБД: PostgreSQL;",
        "обмен файлами: локальное файловое хранилище на сервере;",
        "форматы экспорта: CSV, XLSX, XLS, XLSM;",
        "видеосвязь: WebRTC с использованием STUN-сервера Google.",
    ]:
        add_bullet(document, item)

    add_heading(document, "3. КОНСТРУИРОВАНИЕ")
    add_heading(document, "3.1. Реализация серверной части", level=2)
    add_body_paragraph(
        document,
        "Серверная часть построена модульно. Ключевые доменные зоны разнесены по самостоятельным модулям: auth, courses, "
        "channels, assignments, messages, notifications, audit, video, storage и realtime. Такой разрез оказался удобен, "
        "потому что логика каждого модуля достаточно самостоятельна и может развиваться без жесткой связки с интерфейсом.",
    )
    add_body_paragraph(
        document,
        "Наиболее насыщенным по бизнес-правилам является модуль assignments. Именно в нем сосредоточена логика статусов "
        "задания и сдачи, обработка поздней отправки, уведомления преподавателей, формирование очереди проверки, "
        "экспорт ведомости и история действий по submission.",
    )

    add_heading(document, "3.2. Реализация клиентской части", level=2)
    add_body_paragraph(
        document,
        "Клиентское приложение организовано как набор маршрутов, каждый из которых соответствует рабочему экрану. "
        "Защищенные разделы обернуты в компонент Protected, а текущее состояние пользователя и токена хранится через AuthContext. "
        "Такое решение не перегружает код и позволяет держать правила доступа близко к маршрутам.",
    )
    add_body_paragraph(
        document,
        "Отдельное внимание уделено тем страницам, где пользователь проводит больше всего времени: курсу, заданию и ведомости. "
        "В них реализованы фильтры, пагинация, локальные индикаторы непрочитанного и быстрые действия без лишних переходов. "
        "Это делает интерфейс не просто набором экранов, а действительно рабочим инструментом.",
    )

    add_heading(document, "3.3. Работа в реальном времени", level=2)
    add_body_paragraph(
        document,
        "Для чатов, уведомлений и личного общения по заданию используется WebSocket-шлюз `/ws`. "
        "Сокет-подключение проходит авторизацию по токену, после чего пользователь подключается к персональной комнате и, "
        "при необходимости, к комнатам канала, личного чата или видеосвязи. Такой подход позволил избежать постоянного polling "
        "там, где нужна реакция интерфейса почти сразу после события.",
    )
    add_body_paragraph(
        document,
        "Видеокомната реализована как отдельный сценарий. Сервер фиксирует участников и ограничения по вместимости, "
        "а медиапотоки между клиентами передаются через WebRTC. С практической точки зрения это дало возможность сохранить "
        "в проекте синхронное взаимодействие, не превращая основную LMS в тяжеловесную видеоплатформу.",
    )

    add_heading(document, "3.4. Работа с файлами и экспортом", level=2)
    add_body_paragraph(
        document,
        "В системе выделены разные каталоги хранения для файлов сообщений, материалов задания, студенческих работ и аватаров. "
        "Такое разделение упрощает сопровождение и делает структуру данных более предсказуемой. "
        "При удалении или замене файла сервер синхронно обновляет запись в базе и путь в файловой системе.",
    )
    add_body_paragraph(
        document,
        "Отдельным прикладным преимуществом стали экспортные функции. Администратор может выгрузить курс в CSV, "
        "сформировать журнал проверки, выгрузить дедлайны, а преподаватель или ассистент может получить ведомость в формате Excel. "
        "Это важно не как дополнительная опция, а как мост между веб-приложением и реальной отчетной работой.",
    )

    add_heading(document, "4. ТЕСТИРОВАНИЕ ПРИЛОЖЕНИЯ")
    add_testing_table(document)
    add_heading(document, "4.2. Результаты проверки", level=2)
    add_body_paragraph(
        document,
        "Проведенные сценарии показали, что основные ветки пользовательской работы закрыты и согласованы между собой. "
        "После загрузки файла обновляется история submission, после итоговой сдачи преподаватель видит элемент в очереди проверки, "
        "после проверки студент получает уведомление, а журнал аудита фиксирует само действие и связанные с ним данные.",
    )
    add_body_paragraph(
        document,
        "Самыми чувствительными к ошибкам оказались функции доступа и сценарии реального времени. Именно поэтому акцент в проверке "
        "делался на фильтрацию каналов по группам, ограничение прав на удаление и редактирование сообщений, а также на события "
        "появления новых сообщений и уведомлений без ручного обновления страницы.",
    )

    add_heading(document, "5. ВВОД В ЭКСПЛУАТАЦИЮ")
    add_heading(document, "5.1. Подготовка окружения", level=2)
    add_body_paragraph(
        document,
        "Для локального запуска проекта требуется развернуть PostgreSQL, установить зависимости backend и frontend, "
        "сгенерировать Prisma-клиент и запустить обе части приложения. Архитектура не требует сложной предварительной настройки, "
        "поэтому система удобна как для разработки, так и для демонстрации.",
    )
    add_number(document, "Запустить PostgreSQL через docker compose.")
    add_number(document, "В backend установить зависимости, выполнить `npx prisma generate` и запустить `npm run start:dev`.")
    add_number(document, "Во frontend установить зависимости и запустить `npm run dev`.")
    add_number(document, "Открыть приложение в браузере по адресу локального frontend-сервера.")

    add_heading(document, "5.2. Краткое руководство пользователя", level=2)
    add_body_paragraph(
        document,
        "Работа студента начинается с регистрации и вступления в курс по коду приглашения. После этого ему доступны "
        "те каналы и задания, которые разрешены его роли и его учебной группе. На странице задания студент может загрузить "
        "черновой вариант файла, вернуться к нему позже, а затем отправить итоговую работу на проверку.",
    )
    add_body_paragraph(
        document,
        "Работа преподавателя строится иначе. Он создает курс, настраивает группы, открывает текстовые и assignment-каналы, "
        "добавляет материалы и следит за дедлайнами. Когда студенты начинают отправлять работы, преподаватель использует "
        "очередь проверки, страницу задания, личный чат и журнал действий для сопровождения процесса оценки.",
    )
    add_body_paragraph(
        document,
        "Для администратора важны прежде всего надзорные функции: контроль ролей, экспорт данных и просмотр аудита. "
        "Таким образом, каждый тип пользователя работает в одном приложении, но видит тот набор инструментов, который действительно ему нужен.",
    )
    add_heading(document, "5.3. Экранные формы приложения", level=2)
    add_body_paragraph(
        document,
        "Ниже приведены основные экранные формы системы. Скриншоты иллюстрируют реальный интерфейс проекта и подтверждают, "
        "что в приложении реализованы связанные между собой рабочие страницы для входа, навигации по курсам, общения, проверки работ и контроля уведомлений.",
    )
    add_figure(document, SCREENSHOTS_DIR / "login.png", "Рисунок 4 - Экран авторизации и регистрации", width_cm=10.5)
    add_figure(document, SCREENSHOTS_DIR / "courses.png", "Рисунок 5 - Страница списка курсов", width_cm=16.0)
    add_figure(document, SCREENSHOTS_DIR / "course.png", "Рисунок 6 - Рабочее пространство курса и чат", width_cm=16.0)
    add_figure(document, SCREENSHOTS_DIR / "assignment.png", "Рисунок 7 - Страница задания и проверки работ", width_cm=14.5)
    add_figure(document, SCREENSHOTS_DIR / "gradebook.png", "Рисунок 8 - Ведомость оценок", width_cm=16.0)
    add_figure(document, SCREENSHOTS_DIR / "notifications.png", "Рисунок 9 - Страница уведомлений", width_cm=16.0)
    add_figure(document, SCREENSHOTS_DIR / "review_queue.png", "Рисунок 10 - Очередь проверки", width_cm=16.0)
    add_figure(document, SCREENSHOTS_DIR / "files_library.png", "Рисунок 11 - Библиотека файлов и материалов", width_cm=16.0)

    add_sources_hint(document)
    add_appendix_one(document)
    add_appendix_two(document)


def normalize_table_fonts(document: Document) -> None:
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        run.font.size = Pt(12)


def add_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)


def main() -> None:
    document = Document()
    set_default_font(document)
    configure_page(document)
    fill_document(document)
    normalize_table_fonts(document)
    add_page_numbers(document)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
